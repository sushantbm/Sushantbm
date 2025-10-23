import logging
import re
import time
from typing import Dict, List, Any
from io import BytesIO

import PyPDF2
import docx
import pdfplumber
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

try:
    from langchain_community.embeddings import OpenAIEmbeddings
    from langchain_community.llms import OpenAI
except ImportError:
    OpenAIEmbeddings = None
    OpenAI = None

logger = logging.getLogger(__name__)


class ResumeParser:
    """Utility class for parsing resume files"""

    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'doc', 'txt']

    def extract_text_from_file(self, file_obj, file_type: str) -> str:
        """Extract text content from uploaded file"""
        try:
            if file_type.lower() == 'pdf':
                return self._extract_from_pdf(file_obj)
            elif file_type.lower() in ['docx', 'doc']:
                return self._extract_from_docx(file_obj)
            elif file_type.lower() == 'txt':
                return self._extract_from_txt(file_obj)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            raise

    def _extract_from_pdf(self, file_obj) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            file_obj.seek(0)
            with pdfplumber.open(file_obj) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return self._clean_text(text)
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")

        try:
            file_obj.seek(0)
            pdf_reader = PyPDF2.PdfReader(file_obj)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return self._clean_text(text)
        except Exception as e:
            raise ValueError("Could not extract text from PDF")

    def _extract_from_docx(self, file_obj) -> str:
        """Extract text from DOCX file"""
        try:
            file_obj.seek(0)
            doc = docx.Document(file_obj)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return self._clean_text(text)
        except Exception as e:
            raise ValueError("Could not extract text from DOCX")

    def _extract_from_txt(self, file_obj) -> str:
        """Extract text from TXT file"""
        try:
            file_obj.seek(0)
            content = file_obj.read()
            if isinstance(content, bytes):
                text = content.decode('utf-8', errors='ignore')
            else:
                text = content
            return self._clean_text(text)
        except Exception as e:
            raise ValueError("Could not extract text from TXT")

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        text = re.sub(r'[^\w\s@.-]', ' ', text)
        return text


class ResumeInfoExtractor:
    """Extract structured information from resume"""

    def __init__(self):
        self.skills_keywords = self._load_skills_keywords()

    def extract_contact_info(self, text: str) -> Dict[str, Any]:
        """Extract contact information"""
        contact_info = {}
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        contact_info['emails'] = list(set(emails))

        phone_patterns = [
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
            r'\(\d{3}\)\s*\d{3}[-.]?\d{4}',
            r'\+\d{1,3}[-\.\s]?\d{1,4}[-\.\s]?\d{1,4}[-\.\s]?\d{1,9}',
        ]
        phones = []
        for pattern in phone_patterns:
            phones.extend(re.findall(pattern, text))
        contact_info['phones'] = list(set(phones))
        return contact_info

    def extract_skills(self, text: str) -> List[str]:
        """Extract technical skills"""
        text_lower = text.lower()
        found_skills = []
        for skill in self.skills_keywords:
            if skill.lower() in text_lower:
                found_skills.append(skill)
        return list(set(found_skills))

    def extract_experience_years(self, text: str) -> Dict[str, Any]:
        """Extract experience information"""
        experience_patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'experience[:\s]*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*yrs?\s*(?:of\s*)?experience',
        ]
        years = []
        for pattern in experience_patterns:
            matches = re.findall(pattern, text.lower())
            years.extend([int(match) for match in matches])
        return {'total_years': max(years) if years else 0, 'experience_mentions': years}

    def _load_skills_keywords(self) -> List[str]:
        """Load common technical skills"""
        return [
            'Python', 'Java', 'JavaScript', 'C++', 'C#', 'PHP', 'Ruby', 'Go', 'Rust', 'Swift', 'Kotlin',
            'TypeScript', 'Scala', 'R', 'MATLAB', 'SQL', 'HTML', 'CSS', 'React', 'Angular', 'Vue.js',
            'Node.js', 'Express.js', 'Django', 'Flask', 'FastAPI', 'Spring Boot', 'Laravel', 'jQuery',
            'PostgreSQL', 'MySQL', 'MongoDB', 'Redis', 'SQLite', 'Oracle', 'Cassandra', 'Elasticsearch',
            'AWS', 'Azure', 'Google Cloud', 'Docker', 'Kubernetes', 'Jenkins', 'Git', 'CI/CD', 'Terraform',
            'Machine Learning', 'Deep Learning', 'TensorFlow', 'PyTorch', 'Pandas', 'NumPy', 'Scikit-learn',
            'REST API', 'GraphQL', 'Microservices', 'Agile', 'Scrum', 'Linux'
        ]


class JobFitAnalyzer:
    """Main class for analyzing resume-job fit"""

    def __init__(self, openai_api_key: str = None):
        self.openai_api_key = openai_api_key
        self.embeddings = None
        self.llm = None

        if openai_api_key and OpenAIEmbeddings and OpenAI:
            try:
                self.embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
                self.llm = OpenAI(openai_api_key=openai_api_key, temperature=0.7)
            except Exception as e:
                logger.warning(f"Could not initialize OpenAI: {e}")

        self.resume_parser = ResumeParser()
        self.info_extractor = ResumeInfoExtractor()

    def analyze_resume_job_fit(self, resume_text: str, job_title: str, job_description: str = "") -> Dict[str, Any]:
        """Analyze how well a resume fits a job"""
        start_time = time.time()

        try:
            contact_info = self.info_extractor.extract_contact_info(resume_text)
            skills = self.info_extractor.extract_skills(resume_text)
            experience = self.info_extractor.extract_experience_years(resume_text)

            if self.embeddings:
                similarity_score = self._calculate_embedding_similarity(resume_text, job_title, job_description)
            else:
                similarity_score = self._calculate_tfidf_similarity(resume_text, job_title, job_description)

            if self.llm:
                explanation = self._generate_llm_explanation(resume_text, job_title, job_description, similarity_score, skills)
            else:
                explanation = self._generate_rule_based_explanation(job_title, similarity_score, skills, experience)

            processing_time = time.time() - start_time

            return {
                'similarity_score': round(similarity_score * 100, 1),
                'fit_explanation': explanation,
                'extracted_skills': skills,
                'contact_info': contact_info,
                'experience': experience,
                'processing_time': round(processing_time, 2),
                'analysis_method': 'langchain' if self.embeddings else 'tfidf'
            }
        except Exception as e:
            logger.error(f"Error in analysis: {str(e)}")
            raise

    def _calculate_embedding_similarity(self, resume_text: str, job_title: str, job_description: str) -> float:
        """Calculate similarity using embeddings"""
        try:
            job_text = f"{job_title} {job_description}"
            resume_embedding = self.embeddings.embed_query(resume_text)
            job_embedding = self.embeddings.embed_query(job_text)
            similarity = cosine_similarity([resume_embedding], [job_embedding])[0][0]
            return max(0.0, min(1.0, similarity))
        except Exception as e:
            logger.error(f"Embedding similarity error: {e}")
            return self._calculate_tfidf_similarity(resume_text, job_title, job_description)

    def _calculate_tfidf_similarity(self, resume_text: str, job_title: str, job_description: str) -> float:
        """Calculate similarity using TF-IDF"""
        try:
            job_text = f"{job_title} {job_description}"
            documents = [resume_text, job_text]
            vectorizer = TfidfVectorizer(stop_words='english', ngram_range=(1, 2), max_features=1000)
            tfidf_matrix = vectorizer.fit_transform(documents)
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            return max(0.0, min(1.0, similarity))
        except Exception as e:
            logger.error(f"TF-IDF error: {e}")
            return 0.5

    def _generate_llm_explanation(self, resume_text: str, job_title: str, job_description: str, 
                                 similarity_score: float, skills: List[str]) -> str:
        """Generate explanation using LLM"""
        try:
            prompt = f'''Analyze this resume for "{job_title}".
Job: {job_description}
Resume: {resume_text[:2000]}...
Score: {similarity_score:.2f}
Skills: {', '.join(skills)}

Provide analysis of:
1. Key strengths and relevant experience
2. Skills alignment
3. Areas for improvement
4. Recommendations

Keep response professional.'''
            response = self.llm(prompt)
            return response.strip()
        except Exception as e:
            logger.error(f"LLM generation error: {e}")
            return self._generate_rule_based_explanation(job_title, similarity_score, skills, {})

    def _generate_rule_based_explanation(self, job_title: str, similarity_score: float, 
                                       skills: List[str], experience: Dict[str, Any]) -> str:
        """Generate explanation using rules"""
        score_percentage = similarity_score * 100

        if score_percentage >= 80:
            fit_level = "Excellent"
            recommendation = "Strong alignment with job requirements."
        elif score_percentage >= 60:
            fit_level = "Good"
            recommendation = "Relevant qualifications with room for enhancement."
        elif score_percentage >= 40:
            fit_level = "Moderate"
            recommendation = "Some relevant experience, additional qualifications needed."
        else:
            fit_level = "Limited"
            recommendation = "Limited alignment with job requirements."

        explanation = f'''**Job Fit Analysis for {job_title}**

**Overall Assessment:** {fit_level} match ({score_percentage:.1f}% compatibility)

**Key Findings:**
- Identified {len(skills)} relevant technical skills
- Experience level: {experience.get('total_years', 'Not specified')} years
- Skills found: {', '.join(skills[:10]) if skills else 'Basic skills detected'}

**Recommendation:** {recommendation}

**Next Steps:**
- Review job-specific requirements in detail
- Highlight relevant projects and achievements
- Ensure all key skills are prominently featured
'''
        return explanation.strip()
